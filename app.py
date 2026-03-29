import base64
from html import escape
import os
import re
from io import BytesIO

import requests
import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from dotenv import load_dotenv
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

import styles

load_dotenv()

st.set_page_config(
    page_title="SmartLesson AI",
    layout="wide",
    page_icon="picture/logo.png",
)

AUTH_SERVICE_URL = os.getenv("AUTH_SERVICE_URL", "http://localhost:8001")
LESSON_SERVICE_URL = os.getenv("LESSON_SERVICE_URL", "http://localhost:8002")
AI_SERVICE_URL = os.getenv("AI_SERVICE_URL", "http://localhost:8003")
REQUEST_TIMEOUT = int(os.getenv("SERVICE_TIMEOUT_SECONDS", "30"))

ICONS = {
    "login": "\U0001F511",
    "register": "\U0001FAAA",
    "new_lesson": "\U0001FA84",
    "history": "\U0001F5C2\uFE0F",
    "community": "\U0001F465",
    "account": "\u2699\uFE0F",
    "export": "\u2B07\uFE0F",
    "word": "\U0001F4DD",
    "pdf": "\U0001F4D5",
    "logout": "\U0001F6AA",
    "content": "\U0001F4DA",
    "lesson_item": "\U0001F4D8",
    "info": "\u2139\uFE0F",
    "rename": "\u270F\uFE0F",
    "password": "\U0001F512",
    "upload": "\U0001F4E4",
    "comment": "\U0001F4AC",
    "download": "\U0001F4E5",
    "like": "\U0001F44D",
    "delete": "\U0001F5D1\uFE0F",
    "reply": "\u21A9\uFE0F",
    "avatar": "\U0001F5BC\uFE0F",
    "materials": "\U0001F4E6",
    "reference": "\U0001F4CE",
    "quiz": "\u2753",
    "rubric": "\U0001F4CA",
    "worksheet": "\U0001F4CB",
}


def get_b64(path):
    if os.path.exists(path):
        with open(path, "rb") as file:
            return base64.b64encode(file.read()).decode()
    return ""


def encode_avatar_file(uploaded_file):
    if uploaded_file is None:
        return "", ""

    content_type = uploaded_file.type or ""
    if not content_type.startswith("image/"):
        return None, "Vui lòng chọn file ảnh PNG, JPG hoặc WebP."

    file_bytes = uploaded_file.getvalue()
    if len(file_bytes) > 2 * 1024 * 1024:
        return None, "Ảnh đại diện phải nhỏ hơn 2MB."

    return base64.b64encode(file_bytes).decode(), content_type


def get_avatar_data_uri(avatar_base64, avatar_content_type):
    if avatar_base64 and avatar_content_type:
        return f"data:{avatar_content_type};base64,{avatar_base64}"
    if logo_64:
        return f"data:image/png;base64,{logo_64}"
    return ""


def render_avatar_markup(name, avatar_base64, avatar_content_type, css_class):
    avatar_name = (name or "").strip()
    avatar_letter = escape((avatar_name[:1] or "U").upper())
    avatar_data_uri = ""
    if avatar_base64 and avatar_content_type:
        avatar_data_uri = f"data:{avatar_content_type};base64,{avatar_base64}"

    if avatar_data_uri:
        return f'<img src="{avatar_data_uri}" alt="{escape(avatar_name or "avatar")}" class="{css_class} {css_class}-img">'

    return f'<div class="{css_class}">{avatar_letter}</div>'


def toggle_session_key(key):
    st.session_state[key] = not st.session_state.get(key, False)


def hide_other_reply_forms(active_key):
    for session_key in list(st.session_state.keys()):
        if session_key.startswith("community_reply_form_") and session_key != active_key:
            st.session_state[session_key] = False


def set_reply_target(form_key, target_name):
    hide_other_reply_forms(form_key)
    st.session_state[form_key] = not st.session_state.get(form_key, False)
    st.session_state[f"{form_key}_target_name"] = target_name


def extract_reference_text(uploaded_file):
    if uploaded_file is None:
        return "", None

    file_name = uploaded_file.name.lower()

    try:
        if file_name.endswith((".txt", ".md")):
            return uploaded_file.getvalue().decode("utf-8", errors="ignore").strip()[:12000], None
        if file_name.endswith(".docx"):
            document = Document(BytesIO(uploaded_file.getvalue()))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs).strip()
            return text[:12000], None
    except Exception as error:
        return "", f"Không đọc được tài liệu tham chiếu: {error}"

    return "", "Chỉ hỗ trợ tài liệu tham chiếu định dạng TXT, MD hoặc DOCX."


def strip_markdown(content):
    cleaned_lines = []

    for raw_line in content.splitlines():
        line = raw_line.strip()

        if not line:
            cleaned_lines.append("")
            continue

        line = re.sub(r"^\s{0,3}#{1,6}\s*", "", line)
        line = re.sub(r"^\s{0,3}[-*+]\s+", "- ", line)
        line = re.sub(r"^\s{0,3}\d+\.\s+", "- ", line)
        line = re.sub(r"`{1,3}", "", line)
        line = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", line)
        line = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", line)
        line = re.sub(r"(\*\*|__)(.*?)\1", r"\2", line)
        line = re.sub(r"(\*|_)(.*?)\1", r"\2", line)
        line = re.sub(r"^\*{3,}$", "", line)
        line = re.sub(r"^_{3,}$", "", line)
        line = re.sub(r"^-{3,}$", "", line)
        line = line.replace(">", "")

        cleaned_lines.append(line.strip())

    return "\n".join(cleaned_lines).strip()


def parse_markdown_blocks(content):
    blocks = []

    for raw_line in content.splitlines():
        original = raw_line.rstrip()
        stripped = original.strip()

        if not stripped or re.fullmatch(r"[*_-]{3,}", stripped):
            blocks.append({"type": "blank", "text": ""})
            continue

        heading_match = re.match(r"^\s{0,3}(#{1,6})\s+(.*)$", original)
        if heading_match:
            blocks.append({"type": "heading", "level": len(heading_match.group(1)), "text": strip_markdown(heading_match.group(2))})
            continue

        bullet_match = re.match(r"^(\s*)[-*+]\s+(.*)$", original)
        if bullet_match:
            blocks.append({"type": "bullet", "level": min(len(bullet_match.group(1)) // 2, 3), "text": strip_markdown(bullet_match.group(2))})
            continue

        ordered_match = re.match(r"^(\s*)\d+\.\s+(.*)$", original)
        if ordered_match:
            blocks.append({"type": "bullet", "level": min(len(ordered_match.group(1)) // 2, 3), "text": strip_markdown(ordered_match.group(2))})
            continue

        blocks.append({"type": "paragraph", "text": strip_markdown(stripped)})

    return blocks


def is_major_section(text):
    normalized = strip_markdown(text).strip().lower().rstrip(":")
    keywords = ["mục tiêu", "chuẩn bị", "tiến trình dạy học", "hoạt động dạy học", "khởi động", "hình thành kiến thức", "luyện tập", "vận dụng", "củng cố", "đánh giá"]
    return any(keyword in normalized for keyword in keywords)


def export_word(title, content):
    blocks = parse_markdown_blocks(content)
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(13)
    run_fonts = style.element.rPr.get_or_add_rFonts()
    for tag in ["w:eastAsia", "w:ascii", "w:hAnsi", "w:cs"]:
        run_fonts.set(qn(tag), "Times New Roman")

    title_paragraph = doc.add_heading(title.upper(), 0)
    title_paragraph.paragraph_format.space_after = Pt(14)

    for block in blocks:
        if block["type"] == "blank":
            doc.add_paragraph("")
        elif block["type"] == "heading":
            paragraph = doc.add_heading(block["text"], level=min(block["level"], 4))
            paragraph.paragraph_format.space_before = Pt(14 if is_major_section(block["text"]) else 8)
            paragraph.paragraph_format.space_after = Pt(8 if is_major_section(block["text"]) else 4)
        elif block["type"] == "bullet":
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = Pt(18 * block["level"])
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.add_run(block["text"])
        else:
            paragraph = doc.add_paragraph(block["text"])
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(6 if is_major_section(block["text"]) else 3)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def export_pdf(title, content):
    blocks = parse_markdown_blocks(content)
    output = BytesIO()
    pdf = canvas.Canvas(output, pagesize=A4)
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(100, 800, "GIAO AN: " + title)
    y = 770

    for block in blocks:
        if block["type"] == "blank":
            y -= 10
            continue

        if block["type"] == "heading":
            font_name = "Helvetica-Bold"
            font_size = 15 if block["level"] == 1 else 14 if block["level"] == 2 else 13
            x = 80
            text = block["text"].upper() if block["level"] <= 2 else block["text"]
            line_height = 22
            spacing_before = 12 if is_major_section(block["text"]) else 6
            spacing_after = 8 if is_major_section(block["text"]) else 4
        elif block["type"] == "bullet":
            font_name = "Helvetica"
            font_size = 12
            x = 90 + (block["level"] * 18)
            text = f"• {block['text']}"
            line_height = 18
            spacing_before = 2
            spacing_after = 2
        else:
            font_name = "Helvetica"
            font_size = 12
            x = 80
            text = block["text"]
            line_height = 18
            spacing_before = 2
            spacing_after = 6 if is_major_section(block["text"]) else 4

        y -= spacing_before
        pdf.setFont(font_name, font_size)
        wrapped_lines = []
        current_line = ""
        for word in text.split():
            candidate = word if not current_line else f"{current_line} {word}"
            if pdf.stringWidth(candidate, font_name, font_size) <= 430:
                current_line = candidate
            else:
                wrapped_lines.append(current_line)
                current_line = word
        if current_line:
            wrapped_lines.append(current_line)

        for wrapped_line in wrapped_lines:
            if y < 50:
                pdf.showPage()
                y = 800
                pdf.setFont(font_name, font_size)
            pdf.drawString(x, y, wrapped_line)
            y -= line_height

        y -= spacing_after

    pdf.save()
    return output.getvalue()


def request_service(method, base_url, path, *, json=None, params=None):
    try:
        response = requests.request(method=method, url=f"{base_url}{path}", json=json, params=params, timeout=REQUEST_TIMEOUT)
    except requests.RequestException as error:
        return False, f"Không kết nối được tới dịch vụ {base_url}: {error}"

    try:
        payload = response.json()
    except ValueError:
        payload = {}

    if response.ok:
        return True, payload

    message = payload.get("detail") if isinstance(payload, dict) else None
    return False, message or f"Dịch vụ trả về lỗi HTTP {response.status_code}."


logo_64 = get_b64("picture/logo.png")
styles.apply_custom_css()

if "auth" not in st.session_state:
    st.session_state.auth = False
if "community_open_comment_form" not in st.session_state:
    st.session_state.community_open_comment_form = None


def logout():
    for key in [
        "auth",
        "uid",
        "name",
        "username",
        "avatar_base64",
        "avatar_content_type",
        "last_res",
        "lesson_title",
        "lesson_subject",
        "lesson_grade",
        "reference_summary",
        "support_materials",
    ]:
        if key in st.session_state:
            del st.session_state[key]
    st.session_state.auth = False

if not st.session_state.auth:
    _, col, _ = st.columns([1, 1.3, 1])
    with col:
        styles.render_login_logo(logo_64)
        tab_login, tab_reg = st.tabs([f"{ICONS['login']} Đăng nhập", f"{ICONS['register']} Đăng ký"])

        with tab_login:
            username = st.text_input("Tên đăng nhập", key="login_username")
            password = st.text_input("Mật khẩu", type="password", key="login_password")
            if st.button("XÁC NHẬN ĐĂNG NHẬP", use_container_width=True):
                ok, result = request_service("POST", AUTH_SERVICE_URL, "/login", json={"username": username, "password": password})
                if ok:
                    st.session_state.update(
                        {
                            "auth": True,
                            "uid": result["id"],
                            "name": result["fullname"],
                            "username": result["username"],
                            "avatar_base64": result.get("avatar_base64", ""),
                            "avatar_content_type": result.get("avatar_content_type", ""),
                        }
                    )
                    st.rerun()
                else:
                    st.error(result)

        with tab_reg:
            new_username = st.text_input("Tên tài khoản", key="reg_username")
            full_name = st.text_input("Nhập họ và tên giáo viên", key="reg_fullname")
            new_password = st.text_input("Đặt mật khẩu", type="password", key="reg_password")
            avatar_file = st.file_uploader("Tải ảnh đại diện", type=["png", "jpg", "jpeg", "webp"], key="reg_avatar")
            if avatar_file is not None:
                st.image(avatar_file, width=120)
            if st.button("HOÀN TẤT ĐĂNG KÝ", use_container_width=True):
                avatar_payload, avatar_content_type = encode_avatar_file(avatar_file)
                if avatar_payload is None:
                    st.error(avatar_content_type)
                    st.stop()
                ok, result = request_service(
                    "POST",
                    AUTH_SERVICE_URL,
                    "/register",
                    json={
                        "username": new_username,
                        "fullname": full_name,
                        "password": new_password,
                        "avatar_base64": avatar_payload,
                        "avatar_content_type": avatar_content_type,
                    },
                )
                if ok:
                    st.success("Đăng ký thành công! Xin mời đăng nhập.")
                else:
                    st.error(result)
else:
    with st.sidebar:
        sidebar_avatar_src = get_avatar_data_uri(st.session_state.get("avatar_base64", ""), st.session_state.get("avatar_content_type", ""))
        st.markdown(
            f"""
            <div class="app-sidebar-logo">
                <center><img src="{sidebar_avatar_src}" class="app-sidebar-logo-img"></center>
            </div>
            <h3 class="app-sidebar-name">GV: {st.session_state.name}</h3>
            """,
            unsafe_allow_html=True,
        )

        menu_choice = st.radio(
            "CHỨC NĂNG HỆ THỐNG",
            [
                f"{ICONS['new_lesson']} Soạn giáo án mới",
                f"{ICONS['history']} Lịch sử bài dạy",
                f"{ICONS['community']} Chia sẻ giáo án",
                f"{ICONS['account']} Tài khoản",
            ],
        )

        if "last_res" in st.session_state:
            st.markdown("---")
            st.markdown(f"### {ICONS['export']} XUẤT GIÁO ÁN")
            st.download_button(
                f"{ICONS['word']} Tải bản Word",
                data=export_word(st.session_state.lesson_title, st.session_state.last_res),
                file_name=f"GiaoAn_{st.session_state.lesson_title}.docx",
                use_container_width=True,
            )
            st.download_button(
                f"{ICONS['pdf']} Tải bản PDF",
                data=export_pdf(st.session_state.lesson_title, st.session_state.last_res),
                file_name=f"GiaoAn_{st.session_state.lesson_title}.pdf",
                use_container_width=True,
            )

        st.markdown("---")
        if st.button(f"{ICONS['logout']} ĐĂNG XUẤT", use_container_width=True):
            logout()
            st.rerun()

    if menu_choice == f"{ICONS['new_lesson']} Soạn giáo án mới":
        st.markdown(
            """
            <div class="gemini-banner">
                <h1>Thiết kế bài dạy thông minh</h1>
            </div>
            """,
            unsafe_allow_html=True,
        )

        with st.form("main_lesson_form"):
            col1, col2 = st.columns(2)
            with col1:
                mon = st.selectbox("Chọn môn học", ["Tin học", "Toán học", "Ngữ văn", "Tiếng Anh", "Vật lý", "Hóa học"])
                lop = st.selectbox("Chọn khối lớp", [str(i) for i in range(1, 13)])
            with col2:
                sach = st.selectbox("Bộ sách giáo khoa", ["Cánh Diều", "Kết nối tri thức", "Chân trời sáng tạo"])
                ten = st.text_input("Nhập tên bài dạy")
                tg = st.text_input("Thời lượng (Ví dụ: 2 tiết)")

            note = st.text_area("Yêu cầu sư phạm cụ thể (Phương pháp dạy, hoạt động nhóm...)", height=120)
            reference_file = st.file_uploader(
                f"{ICONS['reference']} Tài liệu tham chiếu cho AI",
                type=["txt", "md", "docx"],
                help="AI sẽ ưu tiên bám sát tài liệu này khi soạn giáo án.",
                key="lesson_reference_file",
            )


            if st.form_submit_button("BẮT ĐẦU SOẠN THẢO BẰNG AI", use_container_width=True):
                if not ten.strip():
                    st.warning("Giáo viên vui lòng điền Tên bài dạy trước khi bắt đầu.")
                else:
                    reference_text, reference_error = extract_reference_text(reference_file)
                    if reference_error:
                        st.error(reference_error)
                        st.stop()
                    with st.spinner("AI đang phân tích chương trình và soạn thảo..."):
                        ok, ai_result = request_service(
                            "POST",
                            AI_SERVICE_URL,
                            "/generate",
                            json={
                                "subject": mon,
                                "grade": lop,
                                "book_series": sach,
                                "lesson_title": ten,
                                "duration": tg,
                                "note": note,
                                "reference_material": reference_text,
                            },
                        )
                        if not ok:
                            st.error(ai_result)
                        else:
                            content = ai_result["content"]
                            st.session_state.update({"last_res": content, "lesson_title": ten, "lesson_subject": mon, "lesson_grade": lop, "reference_summary": reference_file.name if reference_file else "", "support_materials": None})
                            save_ok, save_result = request_service(
                                "POST",
                                LESSON_SERVICE_URL,
                                "/lessons",
                                json={
                                    "user_id": st.session_state.uid,
                                    "subject": mon,
                                    "grade": lop,
                                    "book_series": sach,
                                    "lesson_title": ten,
                                    "duration": tg,
                                    "content_md": content,
                                },
                            )
                            if not save_ok:
                                st.error(save_result)
                            else:
                                st.rerun()

        if "last_res" in st.session_state:
            st.markdown(f"### {ICONS['content']} NỘI DUNG GIÁO ÁN CHI TIẾT")
            st.info(st.session_state.last_res)
            if st.session_state.get("reference_summary"):
                st.caption(f"{ICONS['reference']} Tài liệu tham chiếu: {st.session_state.reference_summary}")

            st.markdown(f"### {ICONS['materials']} BỘ TÀI LIỆU HỖ TRỢ")
            if st.button("TẠO CÂU HỎI, RUBRIC VÀ PHIẾU HỌC TẬP", key="generate_support_materials", use_container_width=True):
                with st.spinner("AI đang tạo bộ tài liệu hỗ trợ..."):
                    ok, materials = request_service(
                        "POST",
                        AI_SERVICE_URL,
                        "/generate-materials",
                        json={
                            "subject": st.session_state.get("lesson_subject", mon),
                            "grade": st.session_state.get("lesson_grade", lop),
                            "lesson_title": st.session_state.lesson_title,
                            "lesson_content": st.session_state.last_res,
                        },
                    )
                    if ok:
                        st.session_state.support_materials = materials
                        st.rerun()
                    else:
                        st.error(materials)

            if st.session_state.get("support_materials"):
                materials_tabs = st.tabs([
                    f"{ICONS['quiz']} Câu hỏi kiểm tra",
                    f"{ICONS['rubric']} Rubric đánh giá",
                    f"{ICONS['worksheet']} Phiếu học tập",
                ])
                with materials_tabs[0]:
                    st.markdown(st.session_state.support_materials.get("questions", "Chưa có nội dung."))
                with materials_tabs[1]:
                    st.markdown(st.session_state.support_materials.get("rubric", "Chưa có nội dung."))
                with materials_tabs[2]:
                    st.markdown(st.session_state.support_materials.get("worksheet", "Chưa có nội dung."))

    elif menu_choice == f"{ICONS['history']} Lịch sử bài dạy":
        st.markdown(f"## {ICONS['history']} DANH SÁCH GIÁO ÁN ĐÃ SOẠN")
        ok, history = request_service("GET", LESSON_SERVICE_URL, f"/lessons/{st.session_state.uid}")
        if not ok:
            st.error(history)
        elif not history:
            st.info("Giáo viên chưa thực hiện soạn giáo án nào trên hệ thống.")
        else:
            for item in history:
                with st.expander(f"{ICONS['lesson_item']} Bài: {item['lesson_title']} (Ngày soạn: {item['created_at']})"):
                    st.markdown(item["content_md"])
                    st.download_button(
                        f"{ICONS['word']} Tải lại bản Word",
                        data=export_word(item["lesson_title"], item["content_md"]),
                        file_name=f"ReDownload_{item['lesson_title']}.docx",
                        key=f"re_dl_{item['id']}",
                    )

    elif menu_choice == f"{ICONS['community']} Chia sẻ giáo án":
        st.markdown(f"## {ICONS['community']} CỘNG ĐỒNG GIÁO ÁN")
        with st.container(border=True):
            st.markdown(f"### {ICONS['upload']} Tạo bài chia sẻ")
            st.caption("Bố cục theo kiểu feed: ô đăng bài nằm trên cùng, danh sách giáo án hiển thị trực tiếp phía dưới.")

            with st.form("shared_lesson_upload_form"):
                share_col1, share_col2 = st.columns(2)
                with share_col1:
                    shared_title = st.text_input("Tiêu đề giáo án")
                    shared_subject = st.text_input("Môn học")
                with share_col2:
                    shared_grade = st.text_input("Khối lớp")
                    shared_file = st.file_uploader("Chọn file giáo án", type=["docx", "pdf", "txt", "md"])
                shared_description = st.text_area("Mô tả ngắn / lưu ý cho đồng nghiệp", height=100)

                if st.form_submit_button("Đăng giáo án", use_container_width=False):
                    if not shared_file:
                        st.warning("Vui lòng chọn file giáo án trước khi đăng.")
                    else:
                        file_bytes = shared_file.getvalue()
                        ok, result = request_service(
                            "POST",
                            LESSON_SERVICE_URL,
                            "/shared-lessons",
                            json={
                                "uploader_id": st.session_state.uid,
                                "uploader_name": st.session_state.name,
                                "title": shared_title,
                                "description": shared_description,
                                "subject": shared_subject,
                                "grade": shared_grade,
                                "file_name": shared_file.name,
                                "content_type": shared_file.type or "application/octet-stream",
                                "file_base64": base64.b64encode(file_bytes).decode(),
                            },
                        )
                        if ok:
                            st.success("Đã đăng giáo án lên cộng đồng.")
                            st.rerun()
                        else:
                            st.error(result)

        st.markdown("---")
        st.markdown(f"### {ICONS['lesson_item']} Giáo án được chia sẻ")
        ok, shared_lessons = request_service("GET", LESSON_SERVICE_URL, "/shared-lessons", params={"viewer_id": st.session_state.uid})

        if not ok:
            st.error(shared_lessons)
        elif not shared_lessons:
            st.info("Chưa có giáo án nào được chia sẻ.")
        else:
            for lesson in shared_lessons:
                with st.container(border=True):
                    uploader_avatar_html = render_avatar_markup(
                        lesson.get("uploader_name", ""),
                        lesson.get("uploader_avatar_base64", ""),
                        lesson.get("uploader_avatar_content_type", ""),
                        "community-feed-avatar",
                    )
                    st.markdown(
                        f"""
                        <div class="community-feed-card">
                            <div class="community-feed-header">
                                {uploader_avatar_html}
                                <div class="community-feed-header-text">
                                    <div class="community-feed-author">{escape(lesson['uploader_name'])}</div>
                                    <div class="community-feed-meta">Đã chia sẻ giáo án • {escape(lesson['created_at'])}</div>
                                </div>
                            </div>
                            <div class="community-feed-body">
                                <div class="community-feed-title">{escape(lesson['title'])}</div>
                                <div class="community-feed-caption">{escape(lesson['description'] or 'Không có mô tả').replace(chr(10), '<br>')}</div>
                                <div class="community-feed-file">Tệp đính kèm: {escape(lesson['file_name'])}</div>
                                <div class="community-feed-tags">
                                    <span class="community-feed-tag">{escape(lesson['subject'] or 'Chưa cập nhật')}</span>
                                    <span class="community-feed-tag">{escape(lesson['grade'] or 'Chưa cập nhật')}</span>
                                </div>
                            </div>
                            <div class="community-feed-stats">
                                <span>👍 {lesson['like_count']} lượt thích</span>
                                <span>💬 {lesson['comment_count']} góp ý</span>
                            </div>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )

                    header_actions_col1, header_actions_col2 = st.columns([12, 1])
                    with header_actions_col2:
                        if lesson["uploader_id"] == st.session_state.uid and hasattr(st, "popover"):
                            with st.popover("⋯"):
                                if st.button("Xóa bài đăng", key=f"delete_shared_{lesson['id']}", use_container_width=False):
                                    delete_ok, delete_result = request_service(
                                        "DELETE",
                                        LESSON_SERVICE_URL,
                                        f"/shared-lessons/{lesson['id']}",
                                        params={"requester_id": st.session_state.uid},
                                    )
                                    if delete_ok:
                                        st.rerun()
                                    else:
                                        st.error(delete_result)

                    action_col1, action_col2, action_col3 = st.columns(3)
                    with action_col1:
                        if st.button("💜" if lesson.get("liked_by_viewer") else "🤍", key=f"like_shared_{lesson['id']}", use_container_width=False):
                            like_ok, like_result = request_service(
                                "POST",
                                LESSON_SERVICE_URL,
                                f"/shared-lessons/{lesson['id']}/likes",
                                json={"user_id": st.session_state.uid, "user_name": st.session_state.name},
                            )
                            if like_ok:
                                st.rerun()
                            else:
                                st.error(like_result)
                    with action_col2:
                        if st.button("💬", key=f"toggle_comment_shared_{lesson['id']}", use_container_width=False):
                            st.session_state.community_open_comment_form = None if st.session_state.community_open_comment_form == lesson["id"] else lesson["id"]
                            st.rerun()
                    with action_col3:
                        download_ok, download_result = request_service("GET", LESSON_SERVICE_URL, f"/shared-lessons/{lesson['id']}/download")
                        if download_ok:
                            st.download_button(
                                "⬇️",
                                data=base64.b64decode(download_result["file_base64"]),
                                file_name=download_result["file_name"],
                                mime=download_result["content_type"],
                                key=f"shared_download_{lesson['id']}",
                                use_container_width=False,
                            )
                        else:
                            st.error(download_result)

                    if st.session_state.community_open_comment_form == lesson["id"]:
                        composer_col1, composer_col2 = st.columns([8, 1.4])
                        with composer_col1:
                            st.text_input(
                                "Nhập bình luận",
                                key=f"comment_input_{lesson['id']}",
                                label_visibility="collapsed",
                                placeholder="Viết bình luận...",
                            )
                        with composer_col2:
                            if st.button("➤", key=f"send_comment_{lesson['id']}", use_container_width=False):
                                comment_ok, comment_result = request_service(
                                    "POST",
                                    LESSON_SERVICE_URL,
                                    f"/shared-lessons/{lesson['id']}/comments",
                                    json={
                                        "commenter_id": st.session_state.uid,
                                        "commenter_name": st.session_state.name,
                                        "comment": st.session_state.get(f"comment_input_{lesson['id']}", ""),
                                    },
                                )
                                if comment_ok:
                                    st.session_state[f"comment_input_{lesson['id']}"] = ""
                                    st.session_state.community_open_comment_form = None
                                    st.rerun()
                                else:
                                    st.error(comment_result)

                    comments_ok, comments_result = request_service(
                        "GET",
                        LESSON_SERVICE_URL,
                        f"/shared-lessons/{lesson['id']}/comments",
                        params={"viewer_id": st.session_state.uid},
                    )
                    if comments_ok and comments_result:
                        top_level_comments = [comment for comment in comments_result if comment["parent_comment_id"] is None]
                        reply_comments = {}
                        for comment in comments_result:
                            if comment["parent_comment_id"] is not None:
                                reply_comments.setdefault(comment["parent_comment_id"], []).append(comment)

                        for comment in top_level_comments:
                            comment_avatar_html = render_avatar_markup(
                                comment.get("commenter_name", ""),
                                comment.get("commenter_avatar_base64", ""),
                                comment.get("commenter_avatar_content_type", ""),
                                "community-comment-avatar",
                            )
                            comment_form_key = f"community_reply_form_{lesson['id']}_{comment['id']}"
                            comment_outer_col1, comment_outer_col2 = st.columns([1, 12])
                            with comment_outer_col1:
                                st.markdown(comment_avatar_html, unsafe_allow_html=True)
                            with comment_outer_col2:
                                st.markdown(
                                    f"""
                                    <div class="community-comment-content">
                                        <div class="community-comment-author">{escape(comment['commenter_name'])}</div>
                                        <div class="community-comment-time">{escape(comment['created_at'])}</div>
                                        <div class="community-comment-body">{escape(comment['comment']).replace(chr(10), '<br>')}</div>
                                    </div>
                                    """,
                                    unsafe_allow_html=True,
                                )
                                comment_action_col1, comment_action_col2, comment_action_col3, comment_action_col4 = st.columns([0.8, 0.8, 0.8, 8])
                                with comment_action_col1:
                                    if st.button("💜" if comment.get("liked_by_viewer") else "🤍", key=f"like_comment_{lesson['id']}_{comment['id']}", use_container_width=False):
                                        like_ok, like_result = request_service(
                                            "POST",
                                            LESSON_SERVICE_URL,
                                            f"/shared-lessons/{lesson['id']}/comments/{comment['id']}/likes",
                                            json={"user_id": st.session_state.uid, "user_name": st.session_state.name},
                                        )
                                        if like_ok:
                                            st.rerun()
                                        else:
                                            st.error(like_result)
                                with comment_action_col2:
                                    st.caption(str(comment.get("like_count", 0)))
                                with comment_action_col3:
                                    if hasattr(st, "popover"):
                                        with st.popover("⋯"):
                                            if st.button("Trả lời", key=f"open_reply_{lesson['id']}_{comment['id']}", use_container_width=False):
                                                set_reply_target(comment_form_key, comment["commenter_name"])
                                                st.rerun()
                                            if comment["commenter_id"] == st.session_state.uid or lesson["uploader_id"] == st.session_state.uid:
                                                if st.button("Xóa", key=f"delete_comment_{lesson['id']}_{comment['id']}", use_container_width=False):
                                                    delete_comment_ok, delete_comment_result = request_service(
                                                        "DELETE",
                                                        LESSON_SERVICE_URL,
                                                        f"/shared-lessons/{lesson['id']}/comments/{comment['id']}",
                                                        params={"requester_id": st.session_state.uid},
                                                    )
                                                    if delete_comment_ok:
                                                        st.rerun()
                                                    else:
                                                        st.error(delete_comment_result)

                                if st.session_state.get(comment_form_key):
                                    reply_col1, reply_col2 = st.columns([8, 1.4])
                                    with reply_col1:
                                        st.text_input("Nhập trả lời", key=f"reply_input_{lesson['id']}_{comment['id']}", label_visibility="collapsed", placeholder=f"Trả lời @{st.session_state.get(f'{comment_form_key}_target_name', comment['commenter_name'])}...")
                                    with reply_col2:
                                        if st.button("➤", key=f"send_reply_{lesson['id']}_{comment['id']}", use_container_width=False):
                                            reply_target = st.session_state.get(f"{comment_form_key}_target_name", comment["commenter_name"])
                                            reply_text = st.session_state.get(f"reply_input_{lesson['id']}_{comment['id']}", "").strip()
                                            payload_comment = f"@{reply_target} {reply_text}".strip()
                                            reply_ok, reply_result = request_service(
                                                "POST",
                                                LESSON_SERVICE_URL,
                                                f"/shared-lessons/{lesson['id']}/comments",
                                                json={
                                                    "commenter_id": st.session_state.uid,
                                                    "commenter_name": st.session_state.name,
                                                    "comment": payload_comment,
                                                    "parent_comment_id": comment["id"],
                                                },
                                            )
                                            if reply_ok:
                                                st.session_state[f"reply_input_{lesson['id']}_{comment['id']}"] = ""
                                                st.session_state[comment_form_key] = False
                                                st.rerun()
                                            else:
                                                st.error(reply_result)

                            for reply in reply_comments.get(comment["id"], []):
                                reply_avatar_html = render_avatar_markup(
                                    reply.get("commenter_name", ""),
                                    reply.get("commenter_avatar_base64", ""),
                                    reply.get("commenter_avatar_content_type", ""),
                                    "community-comment-avatar",
                                )
                                reply_form_key = f"community_reply_form_{lesson['id']}_{reply['id']}"
                                reply_outer_col1, reply_outer_col2 = st.columns([1, 12])
                                with reply_outer_col1:
                                    st.markdown(reply_avatar_html, unsafe_allow_html=True)
                                with reply_outer_col2:
                                    st.markdown(
                                        f"""
                                        <div class="community-comment-content community-comment-content-reply">
                                            <div class="community-comment-author">{escape(reply['commenter_name'])}</div>
                                            <div class="community-comment-time">{escape(reply['created_at'])}</div>
                                            <div class="community-comment-body">{escape(reply['comment']).replace(chr(10), '<br>')}</div>
                                        </div>
                                        """,
                                        unsafe_allow_html=True,
                                    )
                                    reply_action_col1, reply_action_col2, reply_action_col3, reply_action_col4 = st.columns([0.8, 0.8, 0.8, 8])
                                    with reply_action_col1:
                                        if st.button("💜" if reply.get("liked_by_viewer") else "🤍", key=f"like_comment_{lesson['id']}_{reply['id']}", use_container_width=False):
                                            like_ok, like_result = request_service(
                                                "POST",
                                                LESSON_SERVICE_URL,
                                                f"/shared-lessons/{lesson['id']}/comments/{reply['id']}/likes",
                                                json={"user_id": st.session_state.uid, "user_name": st.session_state.name},
                                            )
                                            if like_ok:
                                                st.rerun()
                                            else:
                                                st.error(like_result)
                                    with reply_action_col2:
                                        st.caption(str(reply.get("like_count", 0)))
                                    with reply_action_col3:
                                        if hasattr(st, "popover"):
                                            with st.popover("⋯"):
                                                if st.button("Trả lời", key=f"open_reply_{lesson['id']}_{reply['id']}", use_container_width=False):
                                                    set_reply_target(reply_form_key, reply["commenter_name"])
                                                    st.rerun()
                                                if reply["commenter_id"] == st.session_state.uid or lesson["uploader_id"] == st.session_state.uid:
                                                    if st.button("Xóa", key=f"delete_reply_{lesson['id']}_{reply['id']}", use_container_width=False):
                                                        delete_reply_ok, delete_reply_result = request_service(
                                                            "DELETE",
                                                            LESSON_SERVICE_URL,
                                                            f"/shared-lessons/{lesson['id']}/comments/{reply['id']}",
                                                            params={"requester_id": st.session_state.uid},
                                                        )
                                                        if delete_reply_ok:
                                                            st.rerun()
                                                        else:
                                                            st.error(delete_reply_result)

                                    if st.session_state.get(reply_form_key):
                                        nested_reply_col1, nested_reply_col2 = st.columns([8, 1.4])
                                        with nested_reply_col1:
                                            st.text_input("Nhập trả lời", key=f"reply_input_{lesson['id']}_{reply['id']}", label_visibility="collapsed", placeholder=f"Trả lời @{st.session_state.get(f'{reply_form_key}_target_name', reply['commenter_name'])}...")
                                        with nested_reply_col2:
                                            if st.button("➤", key=f"send_reply_{lesson['id']}_{reply['id']}", use_container_width=False):
                                                reply_target = st.session_state.get(f"{reply_form_key}_target_name", reply["commenter_name"])
                                                reply_text = st.session_state.get(f"reply_input_{lesson['id']}_{reply['id']}", "").strip()
                                                payload_comment = f"@{reply_target} {reply_text}".strip()
                                                nested_reply_ok, nested_reply_result = request_service(
                                                    "POST",
                                                    LESSON_SERVICE_URL,
                                                    f"/shared-lessons/{lesson['id']}/comments",
                                                    json={
                                                        "commenter_id": st.session_state.uid,
                                                        "commenter_name": st.session_state.name,
                                                        "comment": payload_comment,
                                                        "parent_comment_id": comment["id"],
                                                    },
                                                )
                                                if nested_reply_ok:
                                                    st.session_state[f"reply_input_{lesson['id']}_{reply['id']}"] = ""
                                                    st.session_state[reply_form_key] = False
                                                    st.rerun()
                                                else:
                                                    st.error(nested_reply_result)
                    elif not comments_ok:
                        st.error(comments_result)

    else:
        st.markdown(f"## {ICONS['account']} QUẢN LÝ TÀI KHOẢN")
        ok, user_info = request_service("GET", AUTH_SERVICE_URL, f"/users/{st.session_state.uid}")

        if not ok:
            st.error(user_info)
        else:
            st.session_state.avatar_base64 = user_info.get("avatar_base64", "")
            st.session_state.avatar_content_type = user_info.get("avatar_content_type", "")
            avatar_preview_src = get_avatar_data_uri(
                user_info.get("avatar_base64", ""),
                user_info.get("avatar_content_type", ""),
            )

            info_col, action_col = st.columns([1, 1])

            with info_col:
                st.markdown(f"### {ICONS['info']} Thông tin hiện tại")
                if avatar_preview_src:
                    st.markdown(
                        f'''
                        <div class="app-sidebar-logo">
                            <center><img src="{avatar_preview_src}" class="app-sidebar-logo-img"></center>
                        </div>
                        ''',
                        unsafe_allow_html=True,
                    )
                st.text_input("Tên đăng nhập", value=user_info["username"], disabled=True)
                st.text_input("Họ và tên", value=user_info["fullname"], disabled=True)

            with action_col:
                with st.form("update_avatar_form"):
                    st.markdown(f"### {ICONS['avatar']} Cập nhật ảnh đại diện")
                    avatar_update_file = st.file_uploader("Chọn ảnh đại diện mới", type=["png", "jpg", "jpeg", "webp"], key="account_avatar")
                    if st.form_submit_button("CẬP NHẬT ẢNH ĐẠI DIỆN", use_container_width=True):
                        avatar_payload, avatar_content_type = encode_avatar_file(avatar_update_file)
                        if avatar_payload is None:
                            st.error(avatar_content_type)
                        else:
                            ok, result = request_service(
                                "PUT",
                                AUTH_SERVICE_URL,
                                f"/users/{st.session_state.uid}/avatar",
                                json={
                                    "avatar_base64": avatar_payload,
                                    "avatar_content_type": avatar_content_type,
                                },
                            )
                            if ok:
                                st.session_state.avatar_base64 = result.get("avatar_base64", "")
                                st.session_state.avatar_content_type = result.get("avatar_content_type", "")
                                st.success(result["message"])
                                st.rerun()
                            else:
                                st.error(result)

                with st.form("update_fullname_form"):
                    st.markdown(f"### {ICONS['rename']} Đổi họ và tên")
                    new_fullname = st.text_input("Họ và tên mới", value=user_info["fullname"])
                    if st.form_submit_button("CẬP NHẬT HỌ TÊN", use_container_width=True):
                        ok, result = request_service(
                            "PUT",
                            AUTH_SERVICE_URL,
                            f"/users/{st.session_state.uid}/fullname",
                            json={"fullname": new_fullname},
                        )
                        if ok:
                            st.session_state.name = result["fullname"]
                            st.success("Đã cập nhật họ và tên.")
                            st.rerun()
                        else:
                            st.error(result)

                with st.form("update_password_form"):
                    st.markdown(f"### {ICONS['password']} Đổi mật khẩu")
                    current_password = st.text_input("Mật khẩu hiện tại", type="password")
                    new_password = st.text_input("Mật khẩu mới", type="password")
                    confirm_password = st.text_input("Xác nhận mật khẩu mới", type="password")
                    if st.form_submit_button("CẬP NHẬT MẬT KHẨU", use_container_width=True):
                        ok, result = request_service(
                            "PUT",
                            AUTH_SERVICE_URL,
                            f"/users/{st.session_state.uid}/password",
                            json={
                                "current_password": current_password,
                                "new_password": new_password,
                                "confirm_password": confirm_password,
                            },
                        )
                        if ok:
                            st.success(result["message"])
                        else:
                            st.error(result)
