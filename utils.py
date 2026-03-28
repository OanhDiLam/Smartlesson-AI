import base64, os
from docx import Document
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from docx.shared import Pt
from docx.oxml.ns import qn

def get_b64_image(path):
    if os.path.exists(path):
        with open(path, "rb") as f: return base64.b64encode(f.read()).decode()
    return ""

def export_word(title, content):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    r = style.element.rPr.get_or_add_rFonts()
    for tag in ['w:eastAsia', 'w:ascii', 'w:hAnsi', 'w:cs']: r.set(qn(tag), 'Times New Roman')
    doc.add_heading(title.upper(), 0)
    doc.add_paragraph(content)
    bio = BytesIO(); doc.save(bio); return bio.getvalue()

def export_pdf(title, content):
    bio = BytesIO(); p = canvas.Canvas(bio, pagesize=A4)
    p.setFont("Helvetica-Bold", 16); p.drawString(100, 800, f"GIAO AN: {title}")
    p.setFont("Helvetica", 12); y = 770
    for line in content.split('\n'):
        if y < 50: p.showPage(); y = 800
        p.drawString(100, y, line[:80]); y -= 20
    p.save(); return bio.getvalue()